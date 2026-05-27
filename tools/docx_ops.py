import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt

from .file_ops import _validate_path
from .security import require_unlocked


def create_docx(path: str, sections: list[dict], title: Optional[str] = None) -> str:
    """Create a .docx from a list of section dicts.

    Each section is a dict supporting:
      - heading: str         (renders as Heading 1)
      - subheading: str      (renders as Heading 2)
      - paragraph: str       (plain prose; \\n becomes a paragraph break)
      - bullets: list[str]   (one item per entry)
      - numbered: list[str]  (numbered list)
      - quote: str           (italic block quote)
      - page_break: bool     (inserts page break before this section)

    Args:
        path: Output .docx path under the allowed base.
        sections: Ordered list of section dicts.
        title: Optional document title used for metadata.
    """
    locked = require_unlocked()
    if locked:
        return locked
    try:
        if not isinstance(sections, list) or not sections:
            return "[create_docx error: 'sections' must be a non-empty list]"
        out = _validate_path(path)
        if out.suffix.lower() != ".docx":
            return "[create_docx error: path must end with .docx]"
        out.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        if title:
            doc.core_properties.title = title
        for sec in sections:
            if not isinstance(sec, dict):
                return "[create_docx error: each section must be an object]"
            _render_section(doc, sec)
        doc.save(str(out))
        return f"Wrote Word document {out} ({len(sections)} section(s))"
    except Exception as e:
        return f"[create_docx error: {type(e).__name__}: {e}]"


def docx_from_markdown(path: str, markdown: str, title: Optional[str] = None) -> str:
    """Convert a markdown document into a .docx.

    Rules:
      - `#` -> Heading 1, `##` -> Heading 2, `###` -> Heading 3
      - `-` / `*` bullets -> bullet list
      - `1.` etc. numbered -> numbered list
      - `>` -> quote
      - blank line -> paragraph break
      - inline code spans `like this` are kept as-is

    Args:
        path: Output .docx path under the allowed base.
        markdown: Markdown source text.
        title: Optional document title metadata.
    """
    sections = _markdown_to_sections(markdown)
    if not sections:
        return "[docx_from_markdown error: no content found]"
    return create_docx(path, sections, title=title)


def read_docx(path: str, max_chars: int = 60_000) -> str:
    """Extract plain text from a .docx (paragraphs and tables)."""
    try:
        p = _validate_path(path)
        if not p.is_file():
            return f"[read_docx error: not a file: {p}]"
        doc = Document(str(p))
        lines: list[str] = []
        for para in doc.paragraphs:
            t = para.text.rstrip()
            if t:
                lines.append(t)
        for i, table in enumerate(doc.tables, 1):
            lines.append(f"\n[table {i}]")
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                lines.append(" | ".join(cells))
        out = "\n".join(lines)
        if len(out) > max_chars:
            out = out[:max_chars] + f"\n... truncated at {max_chars:,} characters"
        return out or "(no extractable text)"
    except Exception as e:
        return f"[read_docx error: {type(e).__name__}: {e}]"


def docx_info(path: str) -> str:
    """Return a structural summary of an existing .docx (paragraph and table counts)."""
    try:
        p = _validate_path(path)
        if not p.is_file():
            return f"[docx_info error: not a file: {p}]"
        doc = Document(str(p))
        meta = doc.core_properties
        headings = [
            para.text.strip()
            for para in doc.paragraphs
            if (para.style.name or "").startswith("Heading") and para.text.strip()
        ]
        lines = [
            f"path: {p}",
            f"title: {meta.title or '(none)'}",
            f"author: {meta.author or '(none)'}",
            f"paragraphs: {len(doc.paragraphs)}",
            f"tables: {len(doc.tables)}",
            f"headings: {len(headings)}",
        ]
        for h in headings[:30]:
            lines.append(f"  - {h}")
        if len(headings) > 30:
            lines.append(f"  ... +{len(headings) - 30} more")
        return "\n".join(lines)
    except Exception as e:
        return f"[docx_info error: {type(e).__name__}: {e}]"


# ---------- internals ----------

def _render_section(doc: Document, sec: dict) -> None:
    if sec.get("page_break"):
        doc.add_page_break()
    heading = sec.get("heading")
    subheading = sec.get("subheading")
    paragraph = sec.get("paragraph")
    bullets = sec.get("bullets")
    numbered = sec.get("numbered")
    quote = sec.get("quote")

    if heading:
        doc.add_heading(str(heading), level=1)
    if subheading:
        doc.add_heading(str(subheading), level=2)
    if paragraph:
        for chunk in str(paragraph).split("\n"):
            p = doc.add_paragraph(chunk)
            for run in p.runs:
                run.font.size = Pt(11)
    if bullets and isinstance(bullets, list):
        for item in bullets:
            doc.add_paragraph(str(item), style="List Bullet")
    if numbered and isinstance(numbered, list):
        for item in numbered:
            doc.add_paragraph(str(item), style="List Number")
    if quote:
        para = doc.add_paragraph(str(quote))
        para.style = "Intense Quote" if "Intense Quote" in [s.name for s in doc.styles] else para.style
        for run in para.runs:
            run.italic = True


def _markdown_to_sections(md: str) -> list[dict]:
    sections: list[dict] = []
    buf_para: list[str] = []
    buf_bullets: list[str] = []
    buf_numbered: list[str] = []
    pending_heading: Optional[tuple[int, str]] = None

    def flush() -> None:
        nonlocal pending_heading, buf_para, buf_bullets, buf_numbered
        if not (pending_heading or buf_para or buf_bullets or buf_numbered):
            return
        sec: dict = {}
        if pending_heading:
            level, text = pending_heading
            sec["heading" if level <= 1 else "subheading"] = text
        if buf_para:
            sec["paragraph"] = "\n".join(buf_para).strip()
        if buf_bullets:
            sec["bullets"] = list(buf_bullets)
        if buf_numbered:
            sec["numbered"] = list(buf_numbered)
        sections.append(sec)
        pending_heading = None
        buf_para = []
        buf_bullets = []
        buf_numbered = []

    for raw in md.splitlines():
        line = raw.rstrip()
        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            flush()
            pending_heading = (len(h.group(1)), h.group(2).strip())
            continue
        b = re.match(r"^\s*[-*]\s+(.*)$", line)
        if b:
            buf_bullets.append(b.group(1).strip())
            continue
        n = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if n:
            buf_numbered.append(n.group(1).strip())
            continue
        if not line.strip():
            if buf_para:
                buf_para.append("")
            continue
        buf_para.append(line)
    flush()
    return sections
